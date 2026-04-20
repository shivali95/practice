#!/usr/bin/env python3
"""Generate Bhakti Saathi Chat PRD — v2 (no API sections, full faith customisation)."""

from docx import Document
from docx.shared import Pt, RGBColor, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

PURPLE = RGBColor(0x35,0x35,0xF3)
DARK   = RGBColor(0x0C,0x0D,0x10)
MID    = RGBColor(0x19,0x1B,0x1E)
LOW    = RGBColor(0x6B,0x6B,0x6B)
WHITE  = RGBColor(0xFF,0xFF,0xFF)

doc = Document()
sec = doc.sections[0]
sec.page_width  = Inches(8.5)
sec.page_height = Inches(11)
sec.left_margin = sec.right_margin = Inches(1)
sec.top_margin  = sec.bottom_margin = Inches(1)
doc.styles['Normal'].font.name = 'Calibri'
doc.styles['Normal'].font.size = Pt(10.5)
doc.styles['Normal'].font.color.rgb = DARK

# ── helpers ──────────────────────────────────────────────────────────────────

def set_cell_bg(cell, hex_color):
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()
    shd = OxmlElement('w:shd')
    shd.set(qn('w:val'),'clear'); shd.set(qn('w:color'),'auto'); shd.set(qn('w:fill'), hex_color)
    tcPr.append(shd)

def rule(color='3535F3'):
    p = doc.add_paragraph()
    pPr = p._p.get_or_add_pPr()
    pBdr = OxmlElement('w:pBdr')
    b = OxmlElement('w:bottom')
    b.set(qn('w:val'),'single'); b.set(qn('w:sz'),'6')
    b.set(qn('w:space'),'1'); b.set(qn('w:color'), color)
    pBdr.append(b); pPr.append(pBdr)
    p.space_before = Pt(2); p.space_after = Pt(2)

def h1(text):
    p = doc.add_paragraph()
    r = p.add_run(text)
    r.font.name='Calibri'; r.font.size=Pt(16); r.font.bold=True; r.font.color.rgb=PURPLE
    p.space_before=Pt(26); p.space_after=Pt(4)
    rule()

def h2(text):
    p = doc.add_paragraph()
    r = p.add_run(text)
    r.font.name='Calibri'; r.font.size=Pt(12.5); r.font.bold=True; r.font.color.rgb=DARK
    p.space_before=Pt(14); p.space_after=Pt(3)

def h3(text):
    p = doc.add_paragraph()
    r = p.add_run(text)
    r.font.name='Calibri'; r.font.size=Pt(11); r.font.bold=True; r.font.color.rgb=MID
    p.space_before=Pt(10); p.space_after=Pt(2)

def body(text, italic=False, color=None):
    p = doc.add_paragraph()
    r = p.add_run(text)
    r.font.name='Calibri'; r.font.size=Pt(10.5)
    r.font.italic=italic; r.font.color.rgb = color or DARK
    p.space_after=Pt(4)
    return p

def bullet(text, level=0):
    p = doc.add_paragraph(style='List Bullet')
    r = p.add_run(text)
    r.font.name='Calibri'; r.font.size=Pt(10.5); r.font.color.rgb=DARK
    p.paragraph_format.left_indent = Inches(0.25*(level+1))
    p.space_after=Pt(2)

def code_block(text):
    p = doc.add_paragraph()
    p.paragraph_format.left_indent  = Inches(0.3)
    p.paragraph_format.right_indent = Inches(0.3)
    p.space_before=Pt(4); p.space_after=Pt(6)
    pPr = p._p.get_or_add_pPr()
    shd = OxmlElement('w:shd')
    shd.set(qn('w:val'),'clear'); shd.set(qn('w:color'),'auto'); shd.set(qn('w:fill'),'F0F0F6')
    pPr.append(shd)
    r = p.add_run(text)
    r.font.name='Courier New'; r.font.size=Pt(8.5); r.font.color.rgb=RGBColor(0x2B,0x2B,0x6B)

def callout(text, bg='E8E8FC', border='3535F3'):
    t = doc.add_table(rows=1, cols=1); t.alignment=WD_TABLE_ALIGNMENT.LEFT
    cell = t.cell(0,0); set_cell_bg(cell, bg)
    tcPr = cell._tc.get_or_add_tcPr()
    borders = OxmlElement('w:tcBorders')
    for edge in ('top','left','bottom','right'):
        e = OxmlElement(f'w:{edge}')
        e.set(qn('w:val'),'single'); e.set(qn('w:sz'),'6')
        e.set(qn('w:space'),'0'); e.set(qn('w:color'), border)
        borders.append(e)
    tcPr.append(borders)
    tcMar = OxmlElement('w:tcMar')
    for side in ('top','left','bottom','right'):
        m = OxmlElement(f'w:{side}')
        m.set(qn('w:w'),'120'); m.set(qn('w:type'),'dxa')
        tcMar.append(m)
    tcPr.append(tcMar)
    p = cell.paragraphs[0]
    r = p.add_run(text)
    r.font.name='Calibri'; r.font.size=Pt(10); r.font.color.rgb=DARK
    doc.add_paragraph().space_after=Pt(2)

def tbl(headers, rows, widths):
    t = doc.add_table(rows=1+len(rows), cols=len(headers))
    t.alignment=WD_TABLE_ALIGNMENT.LEFT
    hr = t.rows[0]
    for i,h in enumerate(headers):
        c = hr.cells[i]; set_cell_bg(c,'3535F3')
        tcPr = c._tc.get_or_add_tcPr()
        tcMar = OxmlElement('w:tcMar')
        for side in ('top','left','bottom','right'):
            m = OxmlElement(f'w:{side}'); m.set(qn('w:w'),'80'); m.set(qn('w:type'),'dxa'); tcMar.append(m)
        tcPr.append(tcMar)
        p = c.paragraphs[0]; r = p.add_run(h)
        r.font.name='Calibri'; r.font.size=Pt(9.5); r.font.bold=True; r.font.color.rgb=WHITE
        c.width=Inches(widths[i])
    for ri,row in enumerate(rows):
        tr = t.rows[ri+1]
        bg = 'F6F6FF' if ri%2==0 else 'FFFFFF'
        for ci,val in enumerate(row):
            c = tr.cells[ci]; set_cell_bg(c,bg)
            tcPr = c._tc.get_or_add_tcPr()
            tcMar = OxmlElement('w:tcMar')
            for side in ('top','left','bottom','right'):
                m = OxmlElement(f'w:{side}'); m.set(qn('w:w'),'80'); m.set(qn('w:type'),'dxa'); tcMar.append(m)
            tcPr.append(tcMar)
            p = c.paragraphs[0]; r = p.add_run(str(val))
            r.font.name='Calibri'; r.font.size=Pt(9.5); r.font.color.rgb=DARK
            c.width=Inches(widths[ci])
    doc.add_paragraph().space_after=Pt(4)

# ═══════════════════════════════════════════════════════════════════════════════
# COVER
# ═══════════════════════════════════════════════════════════════════════════════
p = doc.add_paragraph(); p.space_before=Pt(48)
r = p.add_run('Bhakti Saathi')
r.font.name='Calibri'; r.font.size=Pt(34); r.font.bold=True; r.font.color.rgb=PURPLE

p2 = doc.add_paragraph()
r2 = p2.add_run('Chat System — Backend Logic & Decision PRD')
r2.font.name='Calibri'; r2.font.size=Pt(17); r2.font.color.rgb=MID

p3 = doc.add_paragraph()
r3 = p3.add_run('Version 2.0  |  March 2026  |  Jio Omni')
r3.font.size=Pt(10); r3.font.color.rgb=RGBColor(0x99,0x99,0xAA)
rule()

body(
    'This document defines the complete chat-layer logic for Bhakti Saathi. '
    'It covers intent detection, context management, orchestration, data models, state design, '
    'edge cases, full faith-based content customisation, and progressive onboarding. '
    'A backend engineer should be able to implement the full system from this document alone '
    'without requiring further product clarification. API definitions are intentionally excluded.'
)

# ═══════════════════════════════════════════════════════════════════════════════
# 1. SYSTEM OVERVIEW
# ═══════════════════════════════════════════════════════════════════════════════
h1('1. Chat System Overview')
body(
    'Bhakti Saathi is a chat-first devotional assistant. The chat interface is the primary product surface. '
    'All features — content, guides, live streams, personalisation — are delivered as structured AI messages '
    'within the conversation thread. The system starts in one of two states:'
)
bullet('ONBOARDING — first session, profile not yet collected')
bullet('CHAT — profile exists, user in free-flowing conversation')
body(
    'The system maintains a user profile (faith, language, region, singer, duration, deity) and a conversation '
    'context (festCtx) that controls contextual routing of ambiguous requests.'
)
callout(
    'CRITICAL: Only Faith and Language are collected upfront. All other profile fields are collected lazily '
    'only when the user accesses the relevant feature for the first time. See Section 13 for full Progressive '
    'Onboarding rules.',
    'FFF3ED','F7AB20')

# ═══════════════════════════════════════════════════════════════════════════════
# 2. INTENT TAXONOMY
# ═══════════════════════════════════════════════════════════════════════════════
h1('2. Intent Taxonomy — Complete List')
body('All intents the system must recognise and route. Rule-based substring matching unless noted.')

groups = [
    ('ONBOARDING',['INTENT_GET_STARTED','INTENT_RESET_ONBOARDING']),
    ('DAILY CONTENT',['INTENT_DAILY_RITUAL','INTENT_CUSTOMIZE_RITUAL','INTENT_BEGIN_RITUAL',
                      'INTENT_MANTRA_OF_DAY','INTENT_PANCHANG','INTENT_ISLAMIC_CALENDAR']),
    ('PUJA / IBADAT',['INTENT_PUJA_VIDHI (context-aware)','INTENT_PUJA_VIDHI_SPECIFIC (deity in query)',
                      'INTENT_SAMAGRI_LIST','INTENT_SAMAGRI_SCAN (camera)']),
    ('FESTIVALS',['INTENT_UPCOMING_FESTIVALS','INTENT_FESTIVAL_HOLI','INTENT_FESTIVAL_NAVRATRI',
                  'INTENT_FESTIVAL_RAM_NAVAMI','INTENT_FESTIVAL_RAMZAN','INTENT_FESTIVAL_EID',
                  'INTENT_FESTIVAL_GENERIC (Baisakhi, Easter, Hola Mohalla...)']),
    ('FESTIVAL SUB-FLOWS (when festCtx is set)',['INTENT_HOLIKA_VIDHI','INTENT_HOLI_SAMAGRI',
        'INTENT_HOLI_MANTRA','INTENT_NAVRATRI_PUJA_STEPS','INTENT_NAVRATRI_MANTRA',
        'INTENT_NAVRATRI_BHAJAN','INTENT_GARBA_EVENTS','INTENT_VRAT_RULES_CONTEXT',
        'INTENT_RAM_NAVAMI_VIDHI','INTENT_RAM_NAVAMI_SAMAGRI']),
    ('VRAT / FASTING',['INTENT_VRAT_SELECTOR (bare "vrat")','INTENT_EKADASHI',
        'INTENT_KARVA_CHAUTH','INTENT_SHIVRATRI','INTENT_SOMVAR_VRAT',
        'INTENT_ROZA_GUIDE (Muslim)','INTENT_UPVAS_GUIDE (Jain)']),
    ('LIVE CONTENT',['INTENT_LIVE_DARSHAN (Hindu/Sikh/Jain/Christian)',
        'INTENT_NAMAAZ_TIMES (Muslim)','INTENT_LIVE_NAMAAZ (Muslim stream)']),
    ('FAITH-SPECIFIC — Muslim',['INTENT_QURAN_VERSE','INTENT_DUA','INTENT_NAMAAZ_GUIDE']),
    ('FAITH-SPECIFIC — Sikh',['INTENT_HUKAMNAMA','INTENT_NITNEM','INTENT_ARDAAS']),
    ('FAITH-SPECIFIC — Jain',['INTENT_NAVKAR_MANTRA','INTENT_SAMAYIK']),
    ('FAITH-SPECIFIC — Christian',['INTENT_BIBLE_VERSE','INTENT_MORNING_PRAYER']),
    ('UTILITY',['INTENT_RAHU_KAAL','INTENT_SETTINGS','INTENT_DEVOTIONAL_FORWARDS','INTENT_VOICE_INPUT']),
    ('FALLBACK',['INTENT_GENERIC']),
]
for gname, intents in groups:
    h3(gname)
    for i in intents: bullet(i)

# ═══════════════════════════════════════════════════════════════════════════════
# 3. FEATURE BREAKDOWN
# ═══════════════════════════════════════════════════════════════════════════════
h1('3. Feature-Wise Breakdown')

# 3.1 Onboarding
h2('3.1  Onboarding Flow')
callout(
    'ONBOARDING = 2 QUESTIONS ONLY: Faith → Language → done. '
    'The user sees the home options grid immediately after. '
    'No region, no singer, no duration, no deity is asked at this stage. Ever.',
    'E8FFE6','25AB21')
body('Triggered on first session (no stored profile). Entry point: INTENT_GET_STARTED on CTA tap.')
h3('Mandatory Onboarding (always, upfront)')
bullet('Step 1: Faith — Hindu / Muslim / Sikh / Jain / Christian')
bullet('Step 2: Language — options change based on faith selected')
bullet('→ Show home options grid immediately. Onboarding is complete.')
body('That is it. Two questions. The user is in the product.')
h3('What Happens Next — Lazy Collection Only')
tbl(
    ['When user first taps…','System asks…','Stores to profile'],
    [
        ['Panchang / Live Darshan / Namaaz Times','Region (North / South / East / West India)','userRegion'],
        ['Daily Ritual','Singer preference','userSinger'],
        ['Daily Ritual (continued)','Ritual duration','userDuration'],
        ['Daily Ritual (continued)','Deity / Ishta Devta','userDeity'],
    ],
    [2.5,2.3,1.7]
)
callout(
    'Singer, Duration, and Deity are NEVER asked unless the user explicitly opens Daily Ritual. '
    'Region is NEVER asked unless the user opens a location-dependent feature. '
    'Every other feature — Mantra, Puja Vidhi, Festivals, Forwards, Vrat — works with just Faith + Language.',
    'FFF3ED','F7AB20')
h3('Decision Tree')
code_block(
    "Session start\n"
    "  profile exists?  YES → CHAT state → showAllOptions()\n"
    "                   NO  → ONBOARDING\n\n"
    "Step 1: askFaith()\n"
    "  → user selects faith → store userFaith\n\n"
    "Step 2: askLanguage()\n"
    "  → user selects language → store userLanguage\n"
    "  → showAllOptions()  ← home grid, 2 taps total\n\n"
    "─── lazy collection triggered later ───────────────────\n\n"
    "User taps Panchang / Live Darshan / Namaaz Times\n"
    "  → region already set? YES → show feature directly\n"
    "                         NO  → askRegion() inline → store → show feature\n\n"
    "User taps Daily Ritual\n"
    "  → ritual profile complete? YES → show ritual card directly\n"
    "                              NO  → askSinger() → askDuration() → askDeity()\n"
    "                                    → store all three → show ritual card"
)
h3('Edge Cases')
bullet('User drops after Step 1 (Faith only) → on re-open, resume at Step 2 (Language)')
bullet('User drops after Step 2 → profile is complete for onboarding; land on home grid on re-open')
bullet('Free text during onboarding → match to closest chip or re-show current step')
bullet('Muslim user reaches Deity step during Daily Ritual setup → show Islamic focus options instead: Fajr Focus / Quran Study / Dhikr / Dua')

# 3.2 Home Options Grid
h2('3.2  Home Options Grid')
body('Shown after mandatory onboarding completes and on "Sabhi Features" chip. All labels are faith-aware — see Section 12.3 for per-faith grids. Only features relevant to the user\'s faith are rendered; irrelevant options are fully omitted.')

# 3.3 Daily Ritual
h2('3.3  Daily Ritual')
h3('Trigger phrases')
body('"daily ritual", "meri ritual", "meri daily", "see my ritual"')
h3('Context Rules')
bullet('Does NOT ask for preferences if userDeity / userSinger / userDuration already stored')
bullet('If profile incomplete → inline collect missing fields FIRST, then render ritual card')
bullet('Defaults: 15 min, Ganesh (Hindu) / Fajr (Muslim) / Nitnem (Sikh)')
h3('Faith-to-Content Mapping')
tbl(
    ['Faith','Card Title','Default Tracks'],
    [
        ['Hindu','Pratah Puja Ritual','Om Jai Jagdish Hare, Ganesh Vandana, Hanuman Chalisa'],
        ['Muslim','Fajr Dhikr Session','Surah Al-Fatiha, Morning Adhkar, Ya Nabi Salam'],
        ['Sikh','Amrit Vela Nitnem','Japji Sahib, Jaap Sahib, Anand Sahib'],
        ['Jain','Pratah Samayik','Navkar Mantra, Samayik Path, Pratikraman'],
        ['Christian','Morning Devotion','Amazing Grace, Morning Prayer, How Great Thou Art'],
    ],
    [1.1,2.1,3.3]
)

# 3.4 Panchang
h2('3.4  Panchang / Islamic Calendar')
h3('Trigger phrases')
body('"panchang", "tithi", "islamic calendar", "hijri", "namaaz time", "rahu kaal", "muhurat"')
callout('System NEVER asks "which calendar?" — userFaith fully determines the branch silently.','E8E8FC','3535F3')
code_block(
    "IF userFaith == Muslim → Islamic Calendar (Hijri, Sehri/Iftar, 5 prayer times)\n"
    "ELSE                   → Hindu Panchang (Tithi, Nakshatra, Rahu Kaal, Sunrise, Muhurat)"
)

# 3.5 Puja Vidhi
h2('3.5  Puja Vidhi / Namaaz Guide / Nitnem Guide')
h3('CRITICAL: Context Resolution Order')
callout(
    'This is the most context-sensitive intent. Resolution must follow the exact order below. '
    'NEVER show a deity selector if festCtx or userDeity can already determine the content.',
    'FFF3ED','F7AB20')
code_block(
    "showPujaVidhi(faith):\n\n"
    "  Step 1: festCtx check\n"
    "    holi       → showHolikaVidhi()\n"
    "    navratri   → showNavratriPujaSteps()\n"
    "    ekadashi   → showEkadashiPujaVidhi()\n"
    "    ram_navami → showRamNavamiVidhi()\n\n"
    "  Step 2: userDeity from profile\n"
    "    IF set → acknowledge + jump to showSpecificPujaVidhi(deityKey)\n"
    "    SKIP selector entirely\n\n"
    "  Step 3: no festCtx, no userDeity\n"
    "    → show deity/guide selector grid\n"
    "    → tap → showSpecificPujaVidhi(selectedKey)"
)
h3('Faith Label Mapping')
tbl(
    ['Faith','Feature Label','Options Shown','Steps'],
    [
        ['Hindu','Puja Vidhi','Ganesh, Shiva, Lakshmi, Vishnu, Durga, Ram, Hanuman, Krishna','6–8 steps each'],
        ['Muslim','Namaaz Guide','Salah (single option)','8 steps (Niyyah → Salam)'],
        ['Sikh','Nitnem Guide','Japji, Jaap, Rehras, Kirtan Sohila','7 steps per path'],
        ['Jain','Pratikraman Vidhi','Samayik, Pratikraman','7 steps'],
        ['Christian','Prayer Guide','Morning, Evening, Intercession','5–7 steps'],
    ],
    [1.1,1.4,2.2,1.8]
)

# 3.6 Mantra / Dua of the Day
h2('3.6  Mantra / Dua / Shabad of the Day')
h3('Context Resolution')
code_block(
    "IF festCtx == 'holi'       → showHoliMantra()\n"
    "IF festCtx == 'navratri'   → showNavratriMantra() (day-specific goddess)\n"
    "IF festCtx == 'ekadashi'   → showEkadashiMantra() (Vishnu Sahasranaam)\n"
    "IF festCtx == 'ram_navami' → showRamMantra()\n"
    "ELSE                       → showMantraOfDay(faith)"
)
tbl(
    ['Faith','Label','Script','Count'],
    [
        ['Hindu','Mantra of the Day','Sanskrit (Devanagari)','108 times'],
        ['Muslim','Dua / Dhikr of the Day','Arabic','100 times after Fajr'],
        ['Sikh','Shabad of the Day','Gurmukhi','Continuous Simran'],
        ['Jain','Sutra of the Day','Prakrit','108 times with mala'],
        ['Christian','Today\'s Verse','English','Memorise and reflect'],
    ],
    [1.2,2.0,1.8,1.5]
)

# 3.7 Live
h2('3.7  Live Darshan / Live Namaaz / Live Kirtan')
callout('userFaith fully determines which live stream branch to show. Never ask "which faith?" or "which temple?" unless the user specifies a temple name in the query.','E8FFE6','25AB21')
code_block(
    "IF userFaith == Muslim    → Live Masjid streams\n"
    "IF userFaith == Sikh      → Harmandir Sahib Kirtan + other Gurdwaras\n"
    "IF userFaith == Jain      → Shikharji Tirth\n"
    "IF userFaith == Christian → Vatican / St Thomas Cathedral\n"
    "ELSE (Hindu)              → Temple darshan schedule by region\n\n"
    "If temple name in query (e.g. 'Vaishno Devi')\n"
    "  → resolve to temple slug → skip list → show stream directly"
)

# 3.8 Festivals
h2('3.8  Upcoming Festivals')
body('Festival data is completely isolated by faith. A user only ever sees their own religion\'s festivals.')
tbl(
    ['Faith','Festivals Shown'],
    [
        ['Hindu','Holi, Navratri, Ram Navami, Diwali, Janmashtami, Maha Shivratri, Ekadashi, Karva Chauth, Teej, Makar Sankranti...'],
        ['Muslim','Ramzan, Eid ul Fitr, Eid ul Adha, Shab-e-Barat, Shab-e-Qadr, Milad-un-Nabi, Muharram'],
        ['Sikh','Baisakhi, Hola Mohalla, Gurpurabs (Guru Nanak, Guru Gobind Singh, others), Diwali (Bandi Chhor Divas)'],
        ['Jain','Mahavir Jayanti, Paryushana (8/10 days), Das Lakshan, Diwali (Mahavir Nirvan), Akshaya Tritiya'],
        ['Christian','Christmas, Easter, Good Friday, Palm Sunday, Ash Wednesday, Pentecost, Advent'],
    ],
    [1.2,5.3]
)
callout(
    'Shared cultural events like Diwali appear for Hindu / Sikh / Jain users only. '
    'Sikhs see the Bandi Chhor Divas framing. Jains see the Mahavir Nirvan framing. '
    'A Muslim or Christian user will NEVER see Diwali.',
    'FFF3ED','F7AB20')

h3('festCtx — Festival Context State')
code_block(
    "Set when:   user opens a festival detail (holi, navratri, ramzan, eid, ekadashi, karva, shivratri, somvar)\n"
    "Used by:    PUJA_VIDHI, MANTRA, VRAT_RULES to route without asking again\n"
    "Cleared:    user navigates to unrelated intent OR taps Back / Sabhi Features\n"
    "            OR a new festCtx is set (replaces old one)"
)

# 3.9 Devotional Forwards
h2('3.9  Devotional Forwards — Aaj ki Story')
body('Trigger phrases: "forward", "devotional forward", "whatsapp", "sandesh", "aaj ke", "aaj ki story"')
callout(
    'Every aspect of devotional forwards — imagery, card title, message text, deities referenced, '
    'greeting style, sign-off — is completely different per faith. '
    'Hindu and Muslim forwards must never share card content. '
    'There is no image of any deity on a Muslim forward.',
    'FFF3ED','F7AB20')
h3('Forward Cards — Faith-by-Faith Content Rules')
tbl(
    ['Faith','Card Categories','Imagery','Greeting','Sign-off'],
    [
        ['Hindu','Radha Krishna, Hanuman, Ganesh, Shiva, Good Morning (Subah), Festival',
         'Deity images, temple photos, flower photography',
         'Jai Shri Krishna / Jai Shree Ram / Pranam',
         'Har Har Mahadev / Jai Jinendra / Radhe Radhe'],
        ['Muslim','Morning Dua, Ramzan, Eid, Islamic Quotes, Good Morning',
         'Mosque silhouettes, geometric Islamic art, calligraphy, nature — NO human/deity images',
         'Assalamu Alaikum / Bismillah',
         'InshaAllah / MashaAllah / JazakAllah Khair'],
        ['Sikh','Waheguru, Gurdwara, Hukamnama card, Gurpurab, Good Morning',
         'Harmandir Sahib, Nishan Sahib, langar, nature',
         'Waheguru Ji Ka Khalsa Waheguru Ji Ki Fateh',
         'Chardi Kala / Sat Sri Akal'],
        ['Jain','Mahavir Jayanti, Paryushana, Navkar, Good Morning',
         'Jain tirth, lotus, ahimsa motifs — no deity idol images in Islamic style',
         'Jai Jinendra',
         'Jai Jinendra / Michhami Dukkadam'],
        ['Christian','Bible Verse, Sunday Blessings, Christmas, Easter, Good Morning',
         'Church, cross, sunrise, nature',
         'God bless you / Praise the Lord',
         'Amen / In His Grace'],
    ],
    [1.0,1.8,2.0,1.4,1.3]
)
h3('Imagery Rules by Faith')
bullet('Hindu: deity images are encouraged — Radha Krishna, Ganesha, etc.')
bullet('Muslim: strictly NO human figures, NO deity images. Use mosque photography, Arabic calligraphy, geometric art, sunrise/nature.')
bullet('Sikh: Gurdwara architecture, Nishan Sahib, community scenes. No individual deity.')
bullet('Jain: lotus, ahimsa hand symbol, Jain temples, Tirthankar motifs (stylised, not photographic deity).')
bullet('Christian: cross, church, sunrise, Bible imagery. No deity icons.')
h3('Message Text Rules by Faith')
bullet('Hindu: Hindi, devotional tone, deity names, blessings ("Ganesh ji aapko sukh de")')
bullet('Muslim: mix of Urdu/Hindi, no deity references, Quranic references OK, dua style ("Allah aapko khush rakhe")')
bullet('Sikh: Punjabi/Hindi mix, references to Waheguru ("Waheguru di kirpa hove")')
bullet('Jain: Hindi, ahimsa/peace focus ("Jivan mein shanti aur ahimsa")')
bullet('Christian: English/Hindi mix, biblical references, hope/faith ("May God\'s grace be with you")')
h3('Two-State Bottom Sheet')
tbl(
    ['State','What User Sees','Actions'],
    [
        ['View (default)','Full-bleed photo 420px, message text overlaid on image, dark gradient','Edit (floating) | Add My Name | Share on WhatsApp'],
        ['Edit mode','Image shrinks to 200px preview, text overlay still live-updates','Textarea → live preview | Apply | Back | Share'],
    ],
    [1.4,3.1,2.0]
)

# ═══════════════════════════════════════════════════════════════════════════════
# 4. VRAT / FASTING — FULL FAITH BREAKDOWN
# ═══════════════════════════════════════════════════════════════════════════════
h1('4. Vrat / Fasting — Faith-Specific Rules')
callout(
    'Fasting rules, allowed foods, rituals, and associated items are completely different per faith. '
    'The system must NEVER show Hindu vrat food rules to a Muslim user, or vice versa.',
    'FFF3ED','F7AB20')

h2('4.1  Hindu Vrat')
h3('Trigger phrases')
body('"vrat", "upvas", "ekadashi", "shivratri", "karva chauth", "somvar", "navratri ke vrat"')
h3('Vrat Types')
tbl(
    ['Vrat','Frequency','Associated Deity','Key Rule'],
    [
        ['Ekadashi','2x per month (Shukla/Krishna Paksha)','Vishnu / Krishna','No rice. Fruits, sabudana, sendha namak only'],
        ['Maha Shivratri','Annual','Shiva','Full day & night fast. Milk, belpatra, water only'],
        ['Somvar Vrat','Every Monday','Shiva','Partial or full fast. Break after evening puja'],
        ['Karva Chauth','Annual (Kartik month)','Shiva/Parvati + moon','Nirjala (no water) from sunrise to moonrise'],
        ['Navratri','8–9 days, twice yearly','Durga / 9 avatars','Satvik food: kuttu atta, sabudana, fruits, dairy. No onion/garlic'],
        ['Ram Navami','Annual','Ram','Fast or partial fast. Panjeeri prasad'],
    ],
    [1.6,2.0,1.7,2.2]
)
h3('Vrat Samagri (Hindu)')
tbl(
    ['Item Category','Items'],
    [
        ['Thali & Vessels','Puja thali, small diyas (lamps), aarti thali with camphor'],
        ['Flowers & Leaves','Marigold, rose petals, tulsi leaves (Ekadashi/Vishnu), belpatra (Shiva/Shivratri), champa'],
        ['Food Offerings','Panjeeri (Ram Navami), modak (Ganesh), panchamrit (milk, curd, honey, sugar, ghee)'],
        ['Ritual Items','Roli, akshat (rice), sindoor, kumkum, janeu (sacred thread), kalash with coconut'],
        ['Incense & Light','Agarbatti (incense sticks), camphor, ghee diya, mustard oil diya'],
    ],
    [1.8,4.7]
)
callout('IMPORTANT: The concept of a "puja thali" with diya, kumkum, sindoor, akshat etc. is exclusively Hindu. This entire category must NOT appear for any other faith.','FFF3ED','F7AB20')

h2('4.2  Muslim Roza (Ramzan Fasting)')
h3('Trigger phrases')
body('"roza", "rozah", "ramzan ka roza", "iftaar", "sehri", "roza kholna"')
h3('Roza Rules')
bullet('Complete fast from Sehri (pre-dawn) end to Iftaar (sunset) — no food, no water, no smoking')
bullet('Sehri: last meal before Fajr adhan. Recommended items: dates, water, roti, sehri foods')
bullet('Iftaar: break fast with dates + water first (Sunnah), then full meal')
bullet('Exceptions: illness, travel, pregnancy, menstruation — qaza (makeup) fast later')
bullet('Fidya and Kaffarah rules for missed fasts')
h3('What to Show')
bullet('Sehri time and Fajr adhan time for user\'s location')
bullet('Iftaar/Maghrib time for user\'s location')
bullet('Roza intention (Niyyah) guide')
bullet('Iftaar dua (Arabic + transliteration)')
bullet('Recommended sehri foods and Iftaar dishes')
h3('What NEVER to Show to a Muslim User')
bullet('No puja thali, no diya, no kumkum, no sindoor, no akshat, no belpatra')
bullet('No mantra, no Sanskrit text, no deity idol imagery')
bullet('No "vrat samagri" list — completely irrelevant concept')
bullet('No "prasad" — use "iftaar food" instead')

h2('4.3  Jain Upvas / Paryushana')
h3('Trigger phrases')
body('"upvas", "paryushana", "das lakshan", "samvatsari", "jain vrat"')
h3('Upvas Rules')
bullet('Jain fasting is called Upvas or Tapasya — different from Hindu vrat')
bullet('Strictest form: Nirjala (no water, no food). Common: Ekasana (one meal a day)')
bullet('No root vegetables ever (potatoes, onions, carrots, garlic, radish — harm micro-organisms)')
bullet('Paryushana: 8 days (Shwetambar) or 10 days (Das Lakshan, Digambar) of intensive fasting')
bullet('Samvatsari: final day — universal forgiveness ritual (Michhami Dukkadam)')
h3('Jain Samagri (Different from Hindu)')
bullet('No thali with kumkum/sindoor/akshat — these are Hindu items')
bullet('Jain puja items: Jal (water), Chandan (sandalwood paste), Pushpa (flowers — no marigold stems touching ground), Dhoop (incense), Deepak, Naivedya (sweets without root veg), Akshat (but used differently)')
bullet('Navkar Mantra recitation with rosary (navkar mala)')
bullet('NO meat, eggs, alcohol, root vegetables in any offering')

h2('4.4  Sikh — No Individual Fasting Tradition')
body('Sikhism does not prescribe individual fasting as a spiritual practice. The Guru Granth Sahib Ji discourages ritual fasting.')
callout('For Sikh users: Do NOT show any vrat/fasting guide. If user asks, respond: "Sikhi mein anushthaan ke liye upvas zaroori nahi — Nitnem, Seva, aur Naam Simran hi asli shakti hai." Show Nitnem guide instead.','E8E8FC','3535F3')

# ═══════════════════════════════════════════════════════════════════════════════
# 5. SAMAGRI — FULL FAITH BREAKDOWN
# ═══════════════════════════════════════════════════════════════════════════════
h1('5. Samagri Lists — Faith-Specific Items')
callout(
    'Samagri is NOT a universal concept. Every item on a Hindu puja samagri list is specific to Hindu ritual. '
    'A Muslim user has no samagri list. A Sikh user has langar offerings and seva items, not a puja samagri. '
    'The system must serve completely different item sets per faith and per occasion.',
    'FFF3ED','F7AB20')

h2('5.1  Hindu Samagri — By Deity')
tbl(
    ['Deity / Occasion','Core Samagri'],
    [
        ['Ganesh Puja','Modak, durva grass (21 blades), red flowers, sindoor, akshat, ghee diya, panchamrit, puja thali, kumkum, supari, paan'],
        ['Shiva Puja / Shivratri','Belpatra (3-leafed), dhatura flower, bhang (optional), milk (abhishek), honey, curd, Ganga jal, rudraksha mala, camphor, vibhuti (holy ash)'],
        ['Lakshmi Puja / Diwali','Lotus flowers, yellow turmeric, kumkum, red cloth, coins, kheel-batase, rice, diya with ghee, sindoor, Lakshmi idol/photo'],
        ['Durga / Navratri','Red chunri (dupatta), red flowers (hibiscus/rose), sindoor, coconut, jaggery, earthen pot (kalash), red bangles, marigold garland'],
        ['Vishnu / Ekadashi','Tulsi leaves (essential), yellow flowers, yellow cloth, panchamrit, janeu, conch (shankh), akshat, tulsi mala'],
        ['Ram Navami','Panjeeri, tulsi, roli, akshat, Ram-Laxman-Sita idol/photo, banana leaves, janeu'],
        ['Holika Dahan','Wood/cow dung cakes (upla) for fire, puja thali, roli, akshat, water, flower garlands, raw coconut, new grain (raw wheat/barley)'],
    ],
    [1.6,4.9]
)
h3('Universal Hindu Samagri (present in all pujas)')
bullet('Puja thali (sacred plate) — EXCLUSIVELY Hindu concept')
bullet('Diyas (clay oil lamps) — lighting essential in every Hindu puja')
bullet('Agarbatti (incense sticks)')
bullet('Camphor (kapoor) for aarti')
bullet('Akshat (unbroken rice grains mixed with turmeric)')
bullet('Roli / Kumkum')
bullet('Kalash (metal/clay pot with water, coconut, mango leaves)')
bullet('Paan leaves and supari')

h2('5.2  Muslim — No Samagri List Concept')
body('The concept of a "samagri list" (ritual object checklist) does not apply to Islamic worship.')
tbl(
    ['Context','What to Show Instead'],
    [
        ['Namaz preparation','Wudu (ablution) guide: clean water, clean clothes, prayer mat (janamaz), direction of Qibla, no music playing'],
        ['Iftaar preparation','Dates (khajoor), water, iftaar food items, Maghrib prayer time, iftaar dua'],
        ['Eid preparation','New/clean clothes, attar (non-alcoholic perfume), Eid namaz time, Fitrana amount'],
        ['Home','Prayer mat (janamaz), tasbeeh (prayer beads for dhikr), Quran with stand'],
    ],
    [1.8,4.7]
)
callout('There is no puja thali, no diya, no kumkum, no sindoor, no idol in Islamic worship. These must NEVER appear in any content shown to a Muslim user.','FFF3ED','F7AB20')

h2('5.3  Sikh — Offerings for Gurdwara / Path')
tbl(
    ['Context','Items'],
    [
        ['Gurdwara visit','Covered head (dupatta/rumal), no shoes, hand washing before entering'],
        ['Langar contribution (seva)','Atta (flour), dal, vegetables, ghee, sugar — for community kitchen'],
        ['Path at home','Gutka Sahib (prayer book), Nitnem book, clean place, covering for head'],
        ['Gurpurab celebration','Karah Prasad (semolina halwa made with equal parts atta, ghee, sugar, water)'],
    ],
    [1.8,4.7]
)

h2('5.4  Jain Samagri — Puja Items')
tbl(
    ['Item','Rule'],
    [
        ['Jal (water)','Fresh, clean — for abhishek of Jain Tirthankara idol'],
        ['Chandan (sandalwood paste)','Ground fresh — applied to idol'],
        ['Pushpa (flowers)','Plucked before sunrise — not from ground. Lotus preferred'],
        ['Dhoop (incense)','Non-animal derived only'],
        ['Deepak (lamp)','Ghee or vegetable oil. NO animal fat'],
        ['Naivedya (food offering)','Vegetarian, no root vegetables, no green leafy veg during paryushana'],
        ['Akshat','Used differently from Hindu tradition — placed at feet of Tirthankara only'],
    ],
    [1.8,4.7]
)
callout('Jain samagri uses NO kumkum (colour powder), NO sindoor, NO marigold stems (considered harmful to micro-organisms). These are exclusively Hindu items.','E8E8FC','3535F3')

# ═══════════════════════════════════════════════════════════════════════════════
# 6. DECISION TREES — KEY LOGIC
# ═══════════════════════════════════════════════════════════════════════════════
h1('6. Decision Trees — Core Logic')

h2('6.1  Intent Router (handleResponse)')
body('Normalised lowercase query string. Rules evaluated top-to-bottom. First match wins.')
code_block(
    " 1.  reset / change faith          → RESET_ONBOARDING\n"
    " 2.  upcoming festival / tyohar    → UPCOMING_FESTIVALS\n"
    " 3.  daily ritual / meri ritual    → DAILY_RITUAL\n"
    " 4.  customise ritual              → CUSTOMIZE_RITUAL\n"
    " 5.  panchang / tithi / hijri      → PANCHANG\n"
    " 6.  holi samagri                  → HOLI_SAMAGRI          ← before generic samagri\n"
    " 7.  samagri                       → SAMAGRI (uses festCtx; faith-aware output)\n"
    " 8.  holika                        → HOLIKA_VIDHI\n"
    " 9.  holi guide / holi festival    → HOLI_GUIDE (festCtx=holi)\n"
    "10.  puja vidhi / puja / pooja     → resolve via festCtx → userDeity → selector\n"
    "11.  mantra                        → resolve via festCtx → MANTRA_OF_DAY\n"
    "12.  navratri                      → NAVRATRI_DETAIL (festCtx=navratri)\n"
    "13.  ramzan / ramadan              → RAMZAN_DETAIL (festCtx=ramzan)\n"
    "14.  eid                           → EID_DETAIL\n"
    "15.  festival / baisakhi / easter  → FESTIVAL_GENERIC\n"
    "16.  live darshan / darshan        → LIVE_DARSHAN(faith)\n"
    "17.  namaaz / namaz                → NAMAAZ_TIMES\n"
    "18.  quran                         → QURAN_VERSE\n"
    "19.  dua                           → DUA\n"
    "20.  hukamnama                     → HUKAMNAMA\n"
    "21.  nitnem                        → NITNEM\n"
    "22.  ardaas                        → ARDAAS\n"
    "23.  navkar / samayik              → NAVKAR_MANTRA\n"
    "24.  bible / verse                 → BIBLE_VERSE\n"
    "25.  prayer                        → MORNING_PRAYER(faith)\n"
    "26.  karva chauth / karva          → KARVA_CHAUTH (festCtx=karva)\n"
    "27.  shivratri                     → SHIVRATRI (festCtx=shivratri)\n"
    "28.  somvar vrat                   → SOMVAR_VRAT (festCtx=somvar)\n"
    "29.  roza / rozah / iftaar         → ROZA_GUIDE\n"
    "30.  vrat (bare)                   → VRAT_SELECTOR (faith-filtered)\n"
    "31.  ekadashi                      → EKADASHI\n"
    "32.  forward / whatsapp / sandesh  → DEVOTIONAL_FORWARDS\n"
    "ELSE                               → GENERIC_RESPONSE(query, faith)"
)

h2('6.2  Puja Vidhi Resolution Tree')
code_block(
    "showPujaVidhi(faith):\n\n"
    "  1. festCtx set?\n"
    "       holi       → showHolikaVidhi()\n"
    "       navratri   → showNavratriPujaSteps()\n"
    "       ekadashi   → showEkadashiPujaVidhi()\n"
    "       ram_navami → showRamNavamiVidhi()\n\n"
    "  2. userDeity set in profile?\n"
    "       YES → 'Aapne {deity} ji ko Ishta Devta chuna hai — seedha unki Puja Vidhi!'\n"
    "             showSpecificPujaVidhi(deityKey)  [skip selector]\n\n"
    "  3. Fallback → show deity selector grid"
)

h2('6.3  Samagri Resolution Tree')
code_block(
    "showSamagri(faith, festCtx, userDeity):\n\n"
    "  IF faith == Muslim → show wudu/namaz prep items [NOT samagri list]\n"
    "  IF faith == Sikh   → show langar seva items or path items\n"
    "  IF faith == Jain   → show Jain puja items (no thali, no kumkum)\n"
    "  ELSE (Hindu):\n"
    "    festCtx == 'holi'       → showHoliSamagri() [holika dahan + rang items]\n"
    "    festCtx == 'navratri'   → showNavratriSamagri() [red items, durga specific]\n"
    "    festCtx == 'ekadashi'   → showEkadashiSamagri() [tulsi-heavy]\n"
    "    festCtx == 'shivratri'  → showShivratriSamagri() [belpatra, milk, dhatura]\n"
    "    userDeity set           → showDeitySamagri(userDeity)\n"
    "    ELSE                    → show universal Hindu samagri + ask which deity"
)

h2('6.4  Vrat Selector — Faith Filtering')
code_block(
    "showVratSelector(faith):\n\n"
    "  IF faith == Muslim  → show Roza guide (Ramzan + Sunnah fasts)\n"
    "  IF faith == Jain    → show Upvas guide (Paryushana, Ekasana)\n"
    "  IF faith == Sikh    → 'Sikhism mein ritual fasting prescribed nahi hai'\n"
    "                         + show Nitnem as alternative\n"
    "  ELSE (Hindu)        → show vrat options: Ekadashi, Shivratri, Somvar,\n"
    "                         Karva Chauth, Navratri, Ram Navami"
)

h2('6.5  Context Switching')
code_block(
    "User: 'Show Holi guide'  → festCtx = 'holi'\n"
    "User: 'Puja vidhi'       → festCtx='holi' → showHolikaVidhi()\n"
    "User: 'Samagri'          → festCtx='holi' → showHoliSamagri()\n"
    "User: 'Mantra'           → festCtx='holi' → showHoliMantra()\n"
    "User: 'Panchang'         → festCtx CLEARED to null\n"
    "User: 'Puja vidhi'       → festCtx=null → check userDeity → selector"
)

# ═══════════════════════════════════════════════════════════════════════════════
# 7. DATA MODELS
# ═══════════════════════════════════════════════════════════════════════════════
h1('7. Data Models')

h2('UserProfile')
code_block(
    "{\n"
    "  userId:          string,\n"
    "  faith:           Hindu | Muslim | Sikh | Jain | Christian,\n"
    "  language:        hi | en | ur | ta | te | mr | sa | pa,\n"
    "  region:          north | south | east | west,   // nullable until collected\n"
    "  singer:          string,                        // nullable until Daily Ritual\n"
    "  duration:        5 | 10 | 15 | 30 | 45,        // nullable until Daily Ritual\n"
    "  deity:           string,                        // nullable until Daily Ritual\n"
    "  profileCreatedAt: ISO8601,\n"
    "  lastActiveAt:    ISO8601\n"
    "}"
)

h2('SessionState')
code_block(
    "{\n"
    "  userId:         string,\n"
    "  sessionId:      string,\n"
    "  isOnboarding:   boolean,\n"
    "  onboardingStep: 0–7,\n"
    "  festCtx:        holi|navratri|ekadashi|ram_navami|ramzan|eid|karva|shivratri|somvar|null,\n"
    "  lastIntent:     string,\n"
    "  lastTemple:     string | null\n"
    "}"
)

h2('ForwardCard')
code_block(
    "{\n"
    "  id:          string,\n"
    "  faith:       Hindu | Muslim | Sikh | Jain | Christian,\n"
    "  category:    Bhakti | Festival | Subah | Vrat | IslamicQuote | Gurpurab | BibleVerse,\n"
    "  imageUrl:    string,      // deity photo (Hindu) / mosque/calligraphy (Muslim) / etc.\n"
    "  fallbackBg:  CSS gradient string,\n"
    "  cardTitle:   string,      // faith-appropriate title\n"
    "  message:     string,      // faith-appropriate text, greeting, sign-off\n"
    "  season:      holi | navratri | ramzan | paryushana | christmas | null\n"
    "}"
)

h2('SamagriItem')
code_block(
    "{\n"
    "  id:         string,\n"
    "  faith:      Hindu | Muslim | Sikh | Jain | Christian,\n"
    "  deity:      string | null,   // null = universal for that faith\n"
    "  festival:   string | null,\n"
    "  name:       string,          // e.g. 'Belpatra', 'Janamaz', 'Chandan'\n"
    "  description: string,\n"
    "  optional:   boolean,\n"
    "  quantity:   string | null    // e.g. '21 blades' (durva), '1 litre' (milk)\n"
    "}"
)

h2('VratGuide')
code_block(
    "{\n"
    "  id:         string,\n"
    "  faith:      Hindu | Muslim | Jain,   // Sikh excluded\n"
    "  name:       string,\n"
    "  localName:  string,         // e.g. 'Roza' not 'Fasting'\n"
    "  nextDate:   ISO8601,\n"
    "  rules:      string[],       // faith-specific rules\n"
    "  allowedFoods: string[],\n"
    "  prohibitedFoods: string[],\n"
    "  pujaGuide:  string | null,  // null for Muslim\n"
    "  dua:        string | null,  // null for Hindu/Jain\n"
    "  niyyah:     string | null   // null for Hindu/Jain\n"
    "}"
)

# ═══════════════════════════════════════════════════════════════════════════════
# 8. STATE & MEMORY
# ═══════════════════════════════════════════════════════════════════════════════
h1('8. State & Memory Design')

h2('Session State (in-memory per conversation)')
tbl(
    ['Key','Type','Reset When'],
    [
        ['isOnboarding','boolean','After Faith + Language collected'],
        ['onboardingStep','0–7','After Faith + Language collected'],
        ['festCtx','string | null','Unrelated intent, Back tap, or new festival opens'],
        ['lastTemple','string | null','Session end'],
        ['currentFwdData','map','Bottom sheet closed'],
    ],
    [1.8,1.3,3.4]
)

h2('Persisted State (user DB)')
tbl(
    ['Key','Collected When','Notes'],
    [
        ['userFaith','Step 1 — always upfront','Permanent until user resets'],
        ['userLanguage','Step 2 — always upfront','Permanent until user resets'],
        ['userRegion','First Panchang / Darshan / Namaaz use','Permanent until user resets'],
        ['userSinger','First Daily Ritual use','Permanent until user resets'],
        ['userDuration','First Daily Ritual use','Permanent until user resets'],
        ['userDeity','First Daily Ritual use','Affects puja vidhi default forever after'],
        ['lastViewedFestival','Auto-tracked','30-day TTL'],
        ['lastViewedTemple','Auto-tracked','7-day TTL'],
    ],
    [1.5,2.3,2.7]
)

# ═══════════════════════════════════════════════════════════════════════════════
# 9. EDGE CASES & FALLBACKS
# ═══════════════════════════════════════════════════════════════════════════════
h1('9. Edge Cases & Fallbacks')

h2('9.1  Intent Ambiguity')
tbl(
    ['Query','Issue','Resolution'],
    [
        ['"puja"','Generic','festCtx → userDeity → show selector'],
        ['"samagri"','Which? For whom?','festCtx → faith branch → deity-specific or universal'],
        ['"mantra"','Which mantra?','festCtx first → daily mantra by faith'],
        ['"vrat"','Which vrat?','VRAT_SELECTOR — faith-filtered list'],
        ['"festival"','Which festival?','Show upcoming festivals for user\'s faith only'],
    ],
    [1.4,1.7,3.4]
)

h2('9.2  Cross-Faith Query Errors')
tbl(
    ['Scenario','Correct Handling'],
    [
        ['Muslim user asks "puja vidhi"','Show Namaaz Guide (map intent to correct faith equivalent)'],
        ['Sikh user asks "vrat"','Explain Sikhism doesn\'t prescribe ritual fasting. Suggest Nitnem.'],
        ['Muslim user asks "samagri"','Show namaz preparation items — NOT puja items'],
        ['Hindu user asks "namaaz time"','If faith is Hindu, this is unusual — confirm intent or show panchang muhurat'],
    ],
    [2.5,4.0]
)

h2('9.3  Missing Context Defaults')
tbl(
    ['Missing Field','Default Behaviour'],
    [
        ['region not set','Use North India defaults (Delhi) until user provides region'],
        ['singer not set','Use faith-default artist for ritual playlist'],
        ['duration not set','Default to 15 min'],
        ['deity not set for puja','Show deity selector (do not guess)'],
    ],
    [2.0,4.5]
)

h2('9.4  Fallback Intent (INTENT_GENERIC)')
code_block(
    "IF query contains devotion/bhakti keyword:\n"
    "  → Contextual AI response + faith-appropriate chips\n\n"
    "IF query is off-topic (food, cricket, etc.):\n"
    "  → 'Main ek devotional assistant hoon — iske baare mein kya madad kar sakta hoon?'\n"
    "  → Chips: Daily Ritual | Panchang | Puja Vidhi (all faith-localised labels)"
)

# ═══════════════════════════════════════════════════════════════════════════════
# 10. ORCHESTRATION
# ═══════════════════════════════════════════════════════════════════════════════
h1('10. Orchestration — Context Switching')
body('The system switches cleanly between any two features. Rules:')
bullet('Only festCtx carries forward within a session')
bullet('userProfile always carries forward — never cleared by navigation')
bullet('No state bleeds between unrelated conversation turns')
bullet('Quick reply chips from old messages remain tappable — trigger fresh intent resolution at that moment')
callout(
    'festCtx does NOT automatically clear after one use. It persists until: '
    '(a) user triggers an unrelated intent, (b) taps Back / Sabhi Features, or (c) a new festCtx is set. '
    'This keeps "puja vidhi" and "mantra" contextual after "show navratri".',
    'E8E8FC','3535F3')

# ═══════════════════════════════════════════════════════════════════════════════
# 11. QUICK REPLY LOGIC
# ═══════════════════════════════════════════════════════════════════════════════
h1('11. Quick Reply / Options Logic')
tbl(
    ['Rule','Detail'],
    [
        ['Show options when','Ambiguity exists and user has not specified enough context'],
        ['Skip options when','Intent fully resolved via festCtx or userDeity'],
        ['Max chips per row','4–5 chips; horizontal scroll for more'],
        ['Chips are contextual','Always reflect the feature just shown — never generic every turn'],
        ['Faith-aware labels','Every chip label resolved from faithStrings[faith][chipKey] — never hardcoded'],
        ['Never repeat onboarding chips','Once faith/language/etc. stored, those prompts never reappear'],
    ],
    [2.5,4.0]
)

# ═══════════════════════════════════════════════════════════════════════════════
# 12. FAITH-BASED CUSTOMISATION
# ═══════════════════════════════════════════════════════════════════════════════
doc.add_page_break()
h1('12. Faith-Based Content Customisation')
callout(
    'CRITICAL: The same intent resolves to completely different content, labels, and features depending on '
    'userFaith. No hardcoded strings anywhere. Every label, chip, card, and AI message tone must be resolved '
    'through a faith-aware lookup. Features irrelevant to a faith are omitted entirely, never greyed out.',
    'FFF3ED','F7AB20')

h2('12.1  Terminology Dictionary')
tbl(
    ['Concept','Hindu','Muslim','Sikh','Jain','Christian'],
    [
        ['Daily practice','Puja / Ritual','Ibadat / Namaaz','Nitnem / Path','Samayik','Devotion / Prayer'],
        ['Place of worship','Mandir','Masjid','Gurdwara','Derasar / Tirth','Church / Chapel'],
        ['Spiritual text','Mantra','Dua / Dhikr','Shabad / Bani','Sutra','Scripture / Verse'],
        ['Sacred book','Geeta / Purana','Quran','Guru Granth Sahib Ji','Agam','Bible'],
        ['Fasting','Vrat','Roza','Not prescribed','Upvas','Fast'],
        ['Calendar','Panchang','Islamic / Hijri','Nanakshahi','Paryushana','Liturgical'],
        ['Greeting','Jai Shree Ram','Assalamu Alaikum','Waheguru Ji Ka Khalsa','Jai Jinendra','Praise the Lord'],
        ['App subtitle','Hindu Devotion','Islamic Devotion','Sikh Devotion','Jain Devotion','Christian Devotion'],
        ['Live content','Live Darshan','Live Namaaz','Live Kirtan','Live Puja','Live Service'],
        ['Step-by-step guide','Puja Vidhi','Namaaz Guide','Nitnem Guide','Pratikraman Vidhi','Prayer Guide'],
        ['Ritual items list','Samagri list','Namaz prep items','Langar seva items','Puja items','Not applicable'],
        ['Morning content','Pratah Vandana','Fajr Dhikr','Amrit Vela Nitnem','Pratah Samayik','Morning Devotion'],
        ['Reciter preference','Gaayak / Singer','Qari / Reciter','Raagi','Shravak','Worship Artist'],
    ],
    [1.4,1.0,1.0,1.0,1.0,1.1]
)

h2('12.2  Feature Availability Matrix')
body('Only render features that apply to userFaith. Omit irrelevant features completely.')
tbl(
    ['Feature','Hindu','Muslim','Sikh','Jain','Christian'],
    [
        ['Puja Vidhi','Yes','—','—','—','—'],
        ['Namaaz Guide','—','Yes','—','—','—'],
        ['Nitnem Guide','—','—','Yes','—','—'],
        ['Pratikraman','—','—','—','Yes','—'],
        ['Morning Devotion','—','—','—','—','Yes'],
        ['Panchang','Yes','—','Partial','—','—'],
        ['Islamic Calendar','—','Yes','—','—','—'],
        ['Namaaz Times','—','Yes','—','—','—'],
        ['Hukamnama','—','—','Yes','—','—'],
        ['Bible Verse','—','—','—','—','Yes'],
        ['Quran Verse','—','Yes','—','—','—'],
        ['Navkar Mantra','—','—','—','Yes','—'],
        ['Mantra of Day','Yes','—','—','—','—'],
        ['Dua of Day','—','Yes','—','—','—'],
        ['Live Darshan (temple)','Yes','—','—','Yes','Yes'],
        ['Live Namaaz','—','Yes','—','—','—'],
        ['Live Kirtan','—','—','Yes','—','—'],
        ['Vrat / Fasting Guide','Yes','Yes (Roza)','—','Yes (Upvas)','—'],
        ['Samagri List','Yes','Partial (prep)','Partial (seva)','Yes (different)','—'],
        ['Devotional Forwards','Yes','Yes','Planned','Planned','Planned'],
    ],
    [2.3,0.8,0.8,0.8,0.8,1.0]
)

h2('12.3  Home Options Grid — Per Faith')
tbl(
    ['Slot','Hindu','Muslim','Sikh','Jain','Christian'],
    [
        ['1','Live Darshan','Live Namaaz','Live Kirtan','Live Puja','Live Service'],
        ['2','Meri Daily Ritual','Meri Daily Ibadat','Mera Daily Nitnem','Mera Samayik','Morning Devotion'],
        ['3','Aaj ka Panchang','Islamic Calendar','Hukamnama','Paryushana','Today\'s Verse'],
        ['4','Aane vale Tyohar','Aane vale Tyohar','Aane vale Gurpurabs','Aane vale Parva','Upcoming Observances'],
        ['5','Aaj ka Mantra','Aaj ki Dua','Aaj ka Shabad','Navkar Mantra','Today\'s Scripture'],
        ['6','Puja Vidhi','Namaaz Guide','Path Guide','Pratikraman Vidhi','Prayer Guide'],
        ['7 (full-width)','Aaj ki Story','Aaj ki Story','Planned','Planned','Planned'],
    ],
    [0.8,1.3,1.3,1.2,1.1,1.3]
)

h2('12.4  AI Tone & Language by Faith')
tbl(
    ['Faith','Tone','Example greeting','Sign-off'],
    [
        ['Hindu','Warm, devotional, Hindi-heavy','Jai Shree Ram! Aaj ka mantra yeh hai...','Har Har Mahadev'],
        ['Muslim','Reverent, clean, bilingual Urdu/Hindi','Assalamu Alaikum! Aaj ki dua...','InshaAllah aapka din mubarak rahe'],
        ['Sikh','Joyful, Punjabi warmth','Waheguru Ji Ka Khalsa! Aaj ka shabad...','Chardi Kala vich raho'],
        ['Jain','Serene, ahimsa-centred, gentle','Jai Jinendra! Aaj ka Navkar path...','Jai Jinendra / Michhami Dukkadam'],
        ['Christian','Warm, hopeful, scripture-grounded','Good morning! Today\'s scripture...','God bless you today / Amen'],
    ],
    [1.1,1.4,2.1,1.9]
)

h2('12.5  Festival Content Isolation')
callout(
    'A Muslim user NEVER sees Hindu festivals. A Hindu user NEVER sees Islamic festivals. '
    'Festival feeds are completely isolated by faith. Shared events (Diwali) shown with faith-specific framing only.',
    'FFF3ED','F7AB20')

# ═══════════════════════════════════════════════════════════════════════════════
# 13. PROGRESSIVE ONBOARDING
# ═══════════════════════════════════════════════════════════════════════════════
h1('13. Progressive Onboarding (Lazy Collection)')
callout(
    'Do not front-load a 6-step onboarding before the user sees any value. '
    'Collect only what is needed, exactly when it is needed. '
    'The user reaches their first feature in 2 taps.',
    'E8FFE6','25AB21')

h2('13.1  What\'s Needed When')
tbl(
    ['Data Point','Needed For','Collected When'],
    [
        ['Faith','Everything — all content, labels, features','Step 1 — always upfront'],
        ['Language','All text responses','Step 2 — always upfront'],
        ['Region','Panchang, Live Darshan, Namaaz Times, Festivals','On first use of any of these'],
        ['Singer / Reciter','Daily Ritual only','On first Daily Ritual access only'],
        ['Duration','Daily Ritual only','On first Daily Ritual access only'],
        ['Deity (Ishta Devta)','Daily Ritual + Puja Vidhi default','On first Daily Ritual access only'],
    ],
    [1.3,2.3,2.9]
)
callout(
    'Singer, Duration, and Deity are ONLY required for Daily Ritual. '
    'Do NOT collect them before the user has specifically requested Daily Ritual. '
    'All other features (Panchang, Puja Vidhi, Live Darshan, Mantra, Festivals) work without them.',
    'FFF3ED','F7AB20')

h2('13.2  Profile Completeness States')
tbl(
    ['State','Fields Set','Features Available'],
    [
        ['MINIMAL','faith + language','Home grid, Mantra/Dua, Festivals, Puja selector, Basic chat'],
        ['LOCATION_READY','+ region','Panchang, Live Darshan/Namaaz, Location-aware festivals'],
        ['RITUAL_READY','+ singer + duration + deity','Fully personalised Daily Ritual, Puja Vidhi auto-skip'],
    ],
    [1.7,2.0,2.8]
)

h2('13.3  Never Re-Ask Rule')
callout(
    'IF a data point is already stored in userProfile → NEVER ask it again in any flow. '
    'Use the stored value silently. This applies to all 6 profile fields. '
    'EXCEPTION: User explicitly taps "Change Faith" or "Reset Profile".',
    'E8E8FC','3535F3')

h2('13.4  Inline Collection Pattern (mid-session lazy collection)')
bullet('Do NOT break the conversation flow or navigate to a new screen')
bullet('Show a single focused AI question + chip options (max 5 chips)')
bullet('User selects → store immediately → proceed to feature — no confirmation screen')
bullet('No "great choice!" interstitial — just collect and continue')
bullet('Small inline tag on feature card: "Teri preferences save ho gayi"')

# ── Footer ─────────────────────────────────────────────────────────────────────
doc.add_paragraph()
rule()
p = doc.add_paragraph()
r = p.add_run(
    'Bhakti Saathi Chat PRD v2.0  |  This document is scoped to chat-layer logic, '
    'decision systems, and faith-aware content rules. API definitions and UI design are out of scope. '
    'A backend engineer can implement the full intent handling, state machine, and edge cases '
    'from this document without further product clarification.')
r.font.size=Pt(9); r.font.italic=True; r.font.color.rgb=LOW
p.alignment=WD_ALIGN_PARAGRAPH.CENTER

output = '/Users/shivali.wason/Documents/devotion/docs/BhaktiSaathi_Chat_PRD.docx'
doc.save(output)
print(f'Saved: {output}')
